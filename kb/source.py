"""Source abstraction for the ingest pipeline.

``Source`` is the seam between the core ingest loop and where content comes
from.  Implement the two-method protocol below to add a new origin — remote
APIs, databases, object storage — without touching ``ingest.py``.

Protocol quick-reference
------------------------
:class:`Source`    — origin that exposes its known keys and yields Documents.
:class:`Document`  — unit of content: identity key, lazy loader, change token,
                     recency timestamp, kind, and the chunk-meta fields.

Adding a new source (no edits to ingest.py required)
-----------------------------------------------------
Here is how a ``NotionSource`` would slot in::

    class NotionSource:
        \"\"\"Stream Notion pages as Documents.\"\"\"

        def __init__(self, api_token: str) -> None:
            self._token = api_token

        def candidate_keys(self) -> set[str]:
            # All page IDs currently visible in the workspace.
            return {page["id"] for page in self._list_pages()}

        def documents(self) -> Iterator[Document]:
            for page in self._list_pages():
                yield Document(
                    key=page["id"],
                    display_name=page["title"],
                    content=lambda p=page: self._fetch_markdown(p["id"]),
                    change_token=(page["last_edited_time"],),
                    since_ts=page["last_edited_unix"],
                    kind="note",
                    path=page["url"],
                    filename=page["title"] + ".md",
                    mtime=page["last_edited_unix"],
                    size=0,
                )

    # The ingest loop in ingest.py is driven by any object that satisfies
    # the Source protocol, so nothing there changes:
    #
    #   _ingest_from_source(NotionSource(token), index_dir=".kb_index")

:class:`FileSource` is the only concrete source shipped today.
"""

from __future__ import annotations

import json
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterator

try:
    from typing import Protocol
except ImportError:  # Python < 3.8
    from typing_extensions import Protocol  # type: ignore[assignment]


# ---------------------------------------------------------------------------
# Filesystem constants — all file-aware policy lives here, not in ingest.py
# ---------------------------------------------------------------------------

# Extensions we treat as ingestible text/code (matched case-insensitively).
_TEXT_SUFFIXES = {
    ".md", ".txt", ".rst",
    ".py", ".js", ".ts", ".tsx", ".jsx",
    ".go", ".rs", ".java", ".rb", ".c", ".cpp", ".h", ".sh",
    ".json", ".yaml", ".yml", ".toml",
    ".html", ".css",
}

# Suffixes that indicate source code (everything else → 'note').
_CODE_SUFFIXES = {
    ".py", ".js", ".ts", ".tsx", ".jsx",
    ".go", ".rs", ".java", ".rb", ".c", ".cpp", ".h", ".sh",
    ".json", ".yaml", ".yml", ".toml",
    ".html", ".css",
}

# Directories that are never worth indexing.
_IGNORE_DIRS = {
    "node_modules", ".git", ".venv", "venv", "__pycache__",
    "dist", "build", ".next", "target", ".idea", ".vscode", "coverage",
    ".mypy_cache", ".pytest_cache", ".tox", ".kb_index",
}

# Files larger than this are skipped (avoids embedding megabyte blobs).
_MAX_FILE_BYTES = 1 * 1024 * 1024  # 1 MB
# Binary documents (PDF/.docx) carry far less text per byte than plain text, so
# allow a larger raw size before extraction.
_MAX_DOC_BYTES = 25 * 1024 * 1024  # 25 MB


# Binary document formats indexed by extracting their text. Always notes.
_DOC_SUFFIXES = {".pdf", ".docx"}

# Every suffix the file walk will pick up.
_INDEXABLE_SUFFIXES = _TEXT_SUFFIXES | _DOC_SUFFIXES


def _kind(suffix: str) -> str:
    return "code" if suffix.lower() in _CODE_SUFFIXES else "note"


def _extract_pdf(p: Path) -> str | None:
    """Extract text from a PDF, or ``None`` if unreadable or pypdf is absent."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return None
    try:
        reader = PdfReader(str(p))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return None
    return text if text.strip() else None


def _extract_docx(p: Path) -> str | None:
    """Extract text from a .docx, or ``None`` if unreadable or python-docx is absent."""
    try:
        from docx import Document as _Docx
    except ImportError:
        return None
    try:
        doc = _Docx(str(p))
        text = "\n".join(par.text for par in doc.paragraphs)
    except Exception:
        return None
    return text if text.strip() else None


def _load_file(p: Path) -> str | None:
    """Read and decode *p*, returning ``None`` for binaries and empty files.

    Decode priority:
    1. UTF-16 BOM (``\\xff\\xfe`` or ``\\xfe\\xff``) — legitimately contains
       NUL bytes, so the NUL heuristic must only apply after ruling this out.
    2. UTF-8 BOM (``\\xef\\xbb\\xbf``) — decoded as ``utf-8-sig``.
    3. NUL byte present (and no BOM) → real binary, skip.
    4. Plain UTF-8 with ``errors='replace'``.
    """
    # Binary document formats need format-specific extraction, not byte decoding.
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix == ".docx":
        return _extract_docx(p)

    try:
        raw = p.read_bytes()
    except Exception:
        return None

    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        try:
            content = raw.decode("utf-16")
        except (UnicodeDecodeError, ValueError):
            return None
    elif raw.startswith(b"\xef\xbb\xbf"):
        content = raw.decode("utf-8-sig", errors="replace")
    elif b"\x00" in raw:
        return None  # real binary
    else:
        content = raw.decode("utf-8", errors="replace")

    return content if content.strip() else None


# ---------------------------------------------------------------------------
# Document — unit of content flowing through the ingest pipeline
# ---------------------------------------------------------------------------

@dataclass
class Document:
    """A single ingestible item produced by a :class:`Source`.

    Attributes:
        key:          Unique identity string used to look up old index chunks
                      (``str(path)`` for files).
        display_name: Human-readable label for logs and status output.
        content:      Zero-argument callable returning the decoded text, or
                      ``None`` when the item should be skipped (binary, empty,
                      too large, network error, …).  Called at most once per
                      ingest run.
        change_token: Opaque equality tuple.  When the stored token equals this
                      value the item is unchanged and its existing vectors are
                      reused without re-embedding.  ``(mtime, size)`` for
                      files; any hashable tuple for other sources.
        since_ts:     Unix timestamp used for ``--since`` recency filtering
                      (``stat.st_mtime`` for files).
        kind:         ``'code'`` or ``'note'``; written into every chunk meta.
        path:         Written verbatim to ``meta["path"]`` for every chunk.
        filename:     Written verbatim to ``meta["filename"]`` for every chunk.
        mtime:        Written verbatim to ``meta["mtime"]`` for every chunk.
        size:         Written verbatim to ``meta["size"]`` for every chunk.
    """

    key: str
    display_name: str
    content: Callable[[], str | None]
    change_token: tuple
    since_ts: float
    kind: str
    # Meta fields written verbatim into each chunk dict
    path: str
    filename: str
    mtime: float
    size: int


# ---------------------------------------------------------------------------
# Source protocol
# ---------------------------------------------------------------------------

class Source(Protocol):
    """Origin that feeds :class:`Document` objects into the ingest pipeline.

    Implement both methods to plug a new content origin into the loop in
    ``ingest.py`` without modifying that file.
    """

    def candidate_keys(self) -> set[str]:
        """Return all keys this source currently exposes.

        Used by the ingest loop to detect documents that were previously
        indexed but have since disappeared.  Must be cheap (no content fetch).
        """
        ...

    def documents(self) -> Iterator[Document]:
        """Yield one :class:`Document` per processable item.

        Items that are definitively unprocessable (unreadable files, auth
        failures, …) may be silently omitted; the ingest loop also handles a
        ``None`` return from ``document.content()``.
        """
        ...


# ---------------------------------------------------------------------------
# FileSource — the only concrete source shipped today
# ---------------------------------------------------------------------------

class FileSource:
    """Walk a local directory tree and yield :class:`Document` objects.

    Reproduces the walk behaviour that was previously inline in ``ingest()``:

    * Prunes :data:`_IGNORE_DIRS` in-place so ``os.walk`` never descends.
    * Keeps only files whose suffix appears in :data:`_TEXT_SUFFIXES`.
    * Returns candidates in deterministic sorted order.
    * Skips files larger than :data:`_MAX_FILE_BYTES`.
    * Decodes content via BOM → UTF-16 → UTF-8-sig → NUL-binary → UTF-8
      and returns ``None`` for undecodable or empty files.

    The directory walk is cached on first use so that :meth:`candidate_keys`
    and :meth:`documents` together cost only one ``os.walk`` call.
    """

    def __init__(self, root: str | Path) -> None:
        self._root = Path(root)
        self._paths: list[Path] | None = None

    def _walk(self) -> list[Path]:
        if self._paths is None:
            collected: list[Path] = []
            for dirpath, dirnames, filenames in os.walk(self._root):
                dirnames[:] = sorted(d for d in dirnames if d not in _IGNORE_DIRS)
                for name in filenames:
                    p = Path(dirpath) / name
                    if p.suffix.lower() in _INDEXABLE_SUFFIXES:
                        collected.append(p)
            self._paths = sorted(collected)
        return self._paths

    def candidate_keys(self) -> set[str]:
        """All file paths (as strings) that pass the walk and suffix filter."""
        return {str(p) for p in self._walk()}

    def documents(self) -> Iterator[Document]:
        """Yield one Document per file that passes the stat and size checks."""
        for p in self._walk():
            try:
                stat = os.stat(p)
                max_bytes = (
                    _MAX_DOC_BYTES
                    if p.suffix.lower() in _DOC_SUFFIXES
                    else _MAX_FILE_BYTES
                )
                if stat.st_size > max_bytes:
                    continue
                mtime = stat.st_mtime
                size = stat.st_size
                path_str = str(p)
                yield Document(
                    key=path_str,
                    display_name=p.name,
                    content=lambda _p=p: _load_file(_p),
                    change_token=(mtime, size),
                    since_ts=mtime,
                    kind=_kind(p.suffix),
                    path=path_str,
                    filename=p.name,
                    mtime=mtime,
                    size=size,
                )
            except Exception:
                continue


# ---------------------------------------------------------------------------
# BookmarkSource - Chrome/Edge "Bookmarks" JSON export
# ---------------------------------------------------------------------------

# WebKit/Chrome epoch offset: seconds between 1601-01-01 and 1970-01-01 (UTC).
_WEBKIT_EPOCH_OFFSET = 11644473600

# Deterministic order and display labels for the three top-level roots.
_BOOKMARK_ROOTS = (
    ("bookmark_bar", "Bookmarks bar"),
    ("other", "Other bookmarks"),
    ("synced", "Mobile bookmarks"),
)


def _webkit_to_unix(value):
    """Convert a Chrome/WebKit timestamp to unix seconds.

    Chrome stores timestamps as a decimal *string* of microseconds since
    1601-01-01 UTC.  Returns ``0.0`` for missing, empty, zero, or otherwise
    unparseable values rather than raising.
    """
    if not value:
        return 0.0
    try:
        micros = int(value)
    except (TypeError, ValueError):
        return 0.0
    if micros <= 0:
        return 0.0
    return micros / 1_000_000 - _WEBKIT_EPOCH_OFFSET


class BookmarkSource:
    """Read a Chrome/Edge ``Bookmarks`` JSON file and yield each bookmark.

    The ``Bookmarks`` file is a pure-JSON tree.  Its top-level ``"roots"``
    object holds ``"bookmark_bar"``, ``"other"`` and (usually) ``"synced"``,
    each a *folder* node::

        {"type": "folder", "name": ..., "children": [...]}

    A *url* node - an actual bookmark - looks like::

        {"type": "url", "name": <title>, "url": <url>, "guid": <id>,
         "date_added": <webkit-microseconds-string>, ...}

    This source walks the three roots in the deterministic order
    ``bookmark_bar`` -> ``other`` -> ``synced`` (children in file order),
    collecting every ``type == "url"`` node together with the folder-name
    path leading to it.  The parsed tree is cached on first use so that
    :meth:`candidate_keys` and :meth:`documents` together cost a single read.

    Every access is defensive: a missing, unreadable, non-JSON, or
    ``roots``-less file yields an empty set / empty iterator without raising,
    mirroring :class:`FileSource`.
    """

    def __init__(self, path: str | Path) -> None:
        self._path = Path(path)
        # Cached list of (url_node, folder_path) tuples; None until walked.
        self._bookmarks: list[tuple[dict, list[str]]] | None = None

    def _walk(self) -> list[tuple[dict, list[str]]]:
        """Parse the file once and return ``(url_node, folder_path)`` pairs.

        Returns an empty list - never raises - when the file is missing,
        unreadable, not valid JSON, or has no ``"roots"`` mapping.
        """
        if self._bookmarks is None:
            collected: list[tuple[dict, list[str]]] = []
            try:
                raw = self._path.read_text(encoding="utf-8")
                data = json.loads(raw)
                roots = data["roots"]
            except Exception:
                self._bookmarks = []
                return self._bookmarks

            for root_key, root_label in _BOOKMARK_ROOTS:
                node = roots.get(root_key) if isinstance(roots, dict) else None
                if isinstance(node, dict):
                    self._collect(node, [root_label], collected)
            self._bookmarks = collected
        return self._bookmarks

    def _collect(
        self,
        node: dict,
        folder_path: list[str],
        out: list[tuple[dict, list[str]]],
    ) -> None:
        """Recursively gather url nodes under *node*, tracking folder names."""
        node_type = node.get("type")
        if node_type == "url":
            out.append((node, folder_path))
            return
        if node_type == "folder":
            children = node.get("children")
            if isinstance(children, list):
                for child in children:
                    if isinstance(child, dict):
                        self._collect(child, folder_path, out)

    @staticmethod
    def _key(node: dict) -> str:
        """Identity key for a url node: prefer ``guid``, else the url."""
        return node.get("guid") or node.get("url") or ""

    def candidate_keys(self) -> set[str]:
        """All bookmark identity keys (guid, falling back to url).

        Cheap: parses the file once (cached) and never fetches external
        content.
        """
        keys = {self._key(node) for node, _ in self._walk()}
        keys.discard("")
        return keys

    def documents(self) -> Iterator[Document]:
        """Yield one :class:`Document` per bookmark (url node)."""
        for node, folder_path in self._walk():
            title = node.get("name") or ""
            url = node.get("url") or ""
            key = self._key(node)
            if not key:
                continue

            date_added = node.get("date_added")
            date_last_used = node.get("date_last_used")
            since_ts = _webkit_to_unix(date_added)
            if since_ts == 0.0:
                since_ts = _webkit_to_unix(date_last_used)

            folder_str = " / ".join(folder_path)

            def _content(_title=title, _url=url, _folder=folder_str):
                if not _title and not _url:
                    return None
                return "\n".join((_title, _url, _folder))

            text = _content()
            size = len(text.encode("utf-8")) if text else 0
            filename = title or url

            yield Document(
                key=key,
                display_name=title,
                content=_content,
                change_token=(date_added, url, title, folder_str),
                since_ts=since_ts,
                kind="note",
                path=url,
                filename=filename,
                mtime=since_ts,
                size=size,
            )
